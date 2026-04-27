"""Geração de memoriais de cálculo em PDF e DOCX.

Cada formato tem sua função geradora que retorna `bytes` prontos para download.
"""
from __future__ import annotations

from datetime import datetime
from io import BytesIO

from . import tabelas as T
from .calculo import ResultadoProjeto
from .modelo import Projeto


# =============================================================================
# Helpers compartilhados
# =============================================================================
def _cabecalho(projeto: Projeto) -> dict:
    e = projeto.estrutura
    return {
        "titulo": "Memorial de Análise de Risco SPDA",
        "subtitulo": "ABNT NBR 5419-2:2026",
        "projeto": e.nome,
        "data": datetime.now().strftime("%d/%m/%Y"),
        "localizacao": f"{e.municipio}/{e.uf}" if e.municipio else "—",
        "ng": f"{e.NG} raios/km²/ano",
    }


def _status_texto(R: float, RT: float) -> str:
    return "NECESSITA PROTEÇÃO" if R > RT else "ACEITÁVEL"


# =============================================================================
# DOCX (Word)
# =============================================================================
def gerar_docx(projeto: Projeto, resultado: ResultadoProjeto) -> bytes:
    """Gera memorial em formato Word (.docx)."""
    from docx import Document
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL

    from . import RT_R1, RT_R3, RT_R4

    doc = Document()
    cab = _cabecalho(projeto)
    e = projeto.estrutura

    # Margens
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # === Cabeçalho ===
    titulo = doc.add_heading(cab["titulo"], level=0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(cab["subtitulo"])
    run.font.size = Pt(12)
    run.italic = True

    # Tabela de identificação
    tab = doc.add_table(rows=4, cols=2)
    tab.style = "Light Grid Accent 1"
    pares = [
        ("Projeto", cab["projeto"]),
        ("Localização", cab["localizacao"]),
        ("NG (Anexo F)", cab["ng"]),
        ("Data de emissão", cab["data"]),
    ]
    for i, (k, v) in enumerate(pares):
        tab.cell(i, 0).text = k
        tab.cell(i, 1).text = str(v)
        for cell in tab.rows[i].cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tab.cell(i, 0).paragraphs[0].runs[0].bold = True

    # === 1. Estrutura ===
    doc.add_heading("1. Características da estrutura", level=1)
    p = doc.add_paragraph()
    p.add_run(f"Dimensões: ").bold = True
    p.add_run(f"L = {e.L} m × W = {e.W} m × H = {e.H} m")
    if e.HP > 0:
        p.add_run(f"  (saliência HP = {e.HP} m)")

    p = doc.add_paragraph()
    p.add_run(f"Localização: ").bold = True
    p.add_run(f"{T.CD_LABELS[e.localizacao.value]} (CD = {T.CD_VALORES[e.localizacao.value]})")

    p = doc.add_paragraph()
    p.add_run(f"Tipo de construção: ").bold = True
    p.add_run(f"{T.RS_LABELS[e.tipo_construcao.value]} (rs = {T.RS_VALORES[e.tipo_construcao.value]})")

    p = doc.add_paragraph()
    p.add_run(f"Total de pessoas (nt): ").bold = True
    p.add_run(f"{e.n_total_pessoas}")

    # === 2. Áreas e eventos ===
    doc.add_heading("2. Áreas equivalentes e eventos perigosos (Anexo A)", level=1)

    tab = doc.add_table(rows=1, cols=3)
    tab.style = "Light Grid Accent 1"
    hdr = tab.rows[0].cells
    hdr[0].text = "Símbolo"
    hdr[1].text = "Valor"
    hdr[2].text = "Equação / Descrição"
    for c in hdr:
        c.paragraphs[0].runs[0].bold = True

    ad_origem = (
        "Informado manualmente (Seção A.2.1.3 — método gráfico)"
        if e.AD_manual is not None and e.AD_manual > 0
        else "Eq. A.1 — área equivalente da estrutura"
    )
    linhas = [
        ("AD", f"{resultado.AD:,.0f} m²", ad_origem),
        ("AM", f"{resultado.AM:,.0f} m²", "Eq. A.6 — descargas próximas (raio 500 m)"),
        ("ND", f"{resultado.ND:.3e}", "Eq. A.3 — descargas na estrutura/ano"),
        ("NM", f"{resultado.NM:.3e}", "Eq. A.5 — descargas próximas/ano"),
        ("NL (total)", f"{resultado.NL_total:.3e}", "Eq. A.7 — descargas em linhas/ano"),
        ("NI (total)", f"{resultado.NI_total:.3e}", "Eq. A.9 — descargas próximas a linhas/ano"),
    ]
    for sim, val, desc in linhas:
        row = tab.add_row().cells
        row[0].text = sim
        row[1].text = val
        row[2].text = desc

    # === 3. Medidas de proteção ===
    doc.add_heading("3. Medidas de proteção", level=1)
    p = doc.add_paragraph()
    p.add_run("SPDA: ").bold = True
    p.add_run(T.PB_LABELS[e.nivel_protecao_spda.value])
    p.add_run(f"  (PB = {T.PB_VALORES[e.nivel_protecao_spda.value]})")

    p = doc.add_paragraph()
    p.add_run("Sistema coordenado de DPS: ").bold = True
    p.add_run(T.PSPD_LABELS[e.pspd_nivel])
    p.add_run(f"  (PSPD = {T.PSPD_VALORES[e.pspd_nivel]})")

    p = doc.add_paragraph()
    p.add_run("Medidas adicionais (PTA): ").bold = True
    p.add_run(", ".join(T.PTA_LABELS[m] for m in e.pta_medidas))

    # === 4. Linhas conectadas ===
    doc.add_heading("4. Linhas elétricas conectadas", level=1)
    if not projeto.linhas:
        doc.add_paragraph("Nenhuma linha externa cadastrada.")
    else:
        tab = doc.add_table(rows=1, cols=5)
        tab.style = "Light Grid Accent 1"
        hdr = tab.rows[0].cells
        for i, h in enumerate(["Nome", "Tipo", "Instalação", "Ambiente", "LL (m) / UW (kV)"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True
        for linha in projeto.linhas:
            row = tab.add_row().cells
            row[0].text = linha.nome
            row[1].text = linha.tipo.value.capitalize()
            row[2].text = T.CI_LABELS[linha.instalacao.value]
            row[3].text = T.CE_LABELS[linha.ambiente.value]
            row[4].text = f"{linha.LL:.0f} / {linha.UW}"

    # === 5. Resultados por zona ===
    doc.add_heading("5. Componentes de risco por zona", level=1)
    for rz in resultado.zonas:
        doc.add_heading(rz.nome, level=2)
        tab = doc.add_table(rows=1, cols=6)
        tab.style = "Light Grid Accent 1"
        hdr = tab.rows[0].cells
        for i, h in enumerate(["Componente", "Descrição", "N", "P", "L", "R"]):
            hdr[i].text = h
            hdr[i].paragraphs[0].runs[0].bold = True

        for nome, c in rz.componentes.items():
            row = tab.add_row().cells
            row[0].text = nome
            row[1].text = c.descricao
            row[2].text = f"{c.N:.3e}"
            row[3].text = f"{c.P:.3e}"
            row[4].text = f"{c.L:.3e}"
            row[5].text = f"{c.R:.3e}"

        p = doc.add_paragraph()
        p.add_run(f"R1 desta zona = {rz.R1:.3e} y⁻¹").bold = True

        # Frequência de danos F
        f = rz.frequencia
        doc.add_heading(f"Frequência de danos F — {rz.nome} (Seção 7)", level=3)
        crit_txt = "crítico (FT = 0,1/ano)" if f.sistema_critico else "não crítico (FT = 1/ano)"
        doc.add_paragraph(f"Classificação: sistema {crit_txt}")

        tabF = doc.add_table(rows=1, cols=2)
        tabF.style = "Light Grid Accent 1"
        hdr = tabF.rows[0].cells
        hdr[0].text = "Componente"
        hdr[1].text = "Valor (1/ano)"
        for c in hdr:
            c.paragraphs[0].runs[0].bold = True
        for nome, val, desc in [
            ("FB", f.FB, "ND × PB (apenas equipamentos em ZPR0A)"),
            ("FC", f.FC, "ND × PC"),
            ("FM", f.FM, "NM × PM"),
            ("FV", f.FV, "(NL + NDJ) × PEB"),
            ("FW", f.FW, "(NL + NDJ) × PW"),
            ("FZ", f.FZ, "NI × PZ"),
        ]:
            row = tabF.add_row().cells
            row[0].text = f"{nome} — {desc}"
            row[1].text = f"{val:.3e}"

        p = doc.add_paragraph()
        run = p.add_run(f"F total = {f.F:.3e}/ano  (FT = {f.FT}/ano)")
        run.bold = True
        from docx.shared import RGBColor
        cor = RGBColor(0xC0, 0x00, 0x00) if f.F > f.FT else RGBColor(0x00, 0x80, 0x00)
        run.font.color.rgb = cor

        status_txt = (
            "Necessita medidas adicionais para reduzir F a FT"
            if f.F > f.FT else "Frequência de danos aceitável"
        )
        doc.add_paragraph(status_txt)

    # === 6. Conclusão ===
    doc.add_heading("6. Riscos totais e conclusão", level=1)
    tab = doc.add_table(rows=1, cols=4)
    tab.style = "Light Grid Accent 1"
    hdr = tab.rows[0].cells
    for i, h in enumerate(["Risco", "Calculado", "Tolerável (RT)", "Status"]):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True

    riscos_linhas = [("R1 — Vida humana", resultado.R1, RT_R1, True)]
    if e.patrimonio_cultural:
        riscos_linhas.append(("R3 — Patrimônio cultural", resultado.R3, RT_R3, True))
    if projeto.avaliar_R4:
        riscos_linhas.append(("R4 — Perda econômica (informativo)", resultado.R4, RT_R4, False))

    for nome, R, RT, normativo in riscos_linhas:
        row = tab.add_row().cells
        row[0].text = nome
        row[1].text = f"{R:.3e}"
        row[2].text = f"{RT:.0e}"
        if normativo:
            row[3].text = _status_texto(R, RT)
            cor = RGBColor(0xC0, 0x00, 0x00) if R > RT else RGBColor(0x00, 0x80, 0x00)
            row[3].paragraphs[0].runs[0].font.color.rgb = cor
            row[3].paragraphs[0].runs[0].bold = True
        else:
            row[3].text = "—"

    # Diagnóstico
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Conclusão: ").bold = True
    if resultado.R1 > RT_R1:
        p.add_run(
            f"O risco R1 calculado ({resultado.R1:.3e}) supera o tolerável "
            f"({RT_R1:.0e}) em {resultado.R1/RT_R1:.1f}×. "
            "Medidas de proteção adicionais devem ser adotadas para reduzir R1 a RT."
        )
    else:
        p.add_run(
            f"O risco R1 calculado ({resultado.R1:.3e}) está abaixo do "
            f"tolerável ({RT_R1:.0e}). A análise atende à NBR 5419-2:2026."
        )

    # Rodapé
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Memorial gerado por Aplicativo de Análise de Risco SPDA — Vertec / NBR 5419-2:2026"
    )
    run.font.size = Pt(9)
    run.italic = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# =============================================================================
# PDF
# =============================================================================
def gerar_pdf(projeto: Projeto, resultado: ResultadoProjeto) -> bytes:
    """Gera memorial em formato PDF."""
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        SimpleDocTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
        PageBreak,
    )

    from . import RT_R1, RT_R3, RT_R4

    cab = _cabecalho(projeto)
    e = projeto.estrutura
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=cab["titulo"],
    )
    estilos = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=estilos["Heading1"],
                        fontSize=14, spaceAfter=8, textColor=colors.HexColor("#1f4e79"))
    h2 = ParagraphStyle("h2", parent=estilos["Heading2"],
                        fontSize=12, spaceAfter=6, textColor=colors.HexColor("#2e74b5"))
    body = ParagraphStyle("body", parent=estilos["BodyText"],
                          fontSize=10, leading=13)
    titulo_style = ParagraphStyle(
        "titulo", parent=estilos["Title"],
        fontSize=18, alignment=1, spaceAfter=4,
    )
    sub_style = ParagraphStyle(
        "sub", parent=estilos["Italic"],
        fontSize=11, alignment=1, spaceAfter=18,
    )

    elementos = []

    def addP(text, style=body):
        elementos.append(Paragraph(text, style))

    def addS(h=6):
        elementos.append(Spacer(1, h))

    def tabela(dados, col_widths=None, header_row=True):
        t = Table(dados, colWidths=col_widths)
        estilo = TableStyle([
            ("FONTNAME", (0, 0), (-1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#888888")),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1),
             [colors.white, colors.HexColor("#f6f6f6")]),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
            ("TOPPADDING", (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ])
        if header_row:
            estilo.add("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79"))
            estilo.add("TEXTCOLOR", (0, 0), (-1, 0), colors.white)
            estilo.add("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold")
        t.setStyle(estilo)
        return t

    # === Cabeçalho ===
    addP(cab["titulo"], titulo_style)
    addP(cab["subtitulo"], sub_style)

    elementos.append(tabela([
        ["Projeto", cab["projeto"]],
        ["Localização", cab["localizacao"]],
        ["NG (Anexo F)", cab["ng"]],
        ["Data de emissão", cab["data"]],
    ], col_widths=[5 * cm, 11 * cm], header_row=False))
    addS(14)

    # === 1. Estrutura ===
    addP("1. Características da estrutura", h1)
    addP(f"<b>Dimensões:</b> L = {e.L} m × W = {e.W} m × H = {e.H} m" +
         (f"  (saliência HP = {e.HP} m)" if e.HP > 0 else ""))
    addP(f"<b>Localização:</b> {T.CD_LABELS[e.localizacao.value]} "
         f"(CD = {T.CD_VALORES[e.localizacao.value]})")
    addP(f"<b>Tipo de construção:</b> {T.RS_LABELS[e.tipo_construcao.value]} "
         f"(rs = {T.RS_VALORES[e.tipo_construcao.value]})")
    addP(f"<b>Total de pessoas (nt):</b> {e.n_total_pessoas}")
    addS(10)

    # === 2. Áreas e eventos ===
    addP("2. Áreas equivalentes e eventos perigosos (Anexo A)", h1)
    ad_origem_pdf = (
        "Informado manualmente (Seção A.2.1.3)"
        if e.AD_manual is not None and e.AD_manual > 0
        else "Eq. A.1 — área equivalente da estrutura"
    )
    elementos.append(tabela([
        ["Símbolo", "Valor", "Equação / Descrição"],
        ["AD", f"{resultado.AD:,.0f} m²", ad_origem_pdf],
        ["AM", f"{resultado.AM:,.0f} m²", "Eq. A.6 — descargas próximas (raio 500 m)"],
        ["ND", f"{resultado.ND:.3e}", "Eq. A.3 — descargas na estrutura/ano"],
        ["NM", f"{resultado.NM:.3e}", "Eq. A.5 — descargas próximas/ano"],
        ["NL (total)", f"{resultado.NL_total:.3e}", "Eq. A.7 — descargas em linhas/ano"],
        ["NI (total)", f"{resultado.NI_total:.3e}", "Eq. A.9 — descargas próximas a linhas/ano"],
    ], col_widths=[3 * cm, 3.5 * cm, 9.5 * cm]))
    addS(10)

    # === 3. Medidas ===
    addP("3. Medidas de proteção", h1)
    addP(f"<b>SPDA:</b> {T.PB_LABELS[e.nivel_protecao_spda.value]} "
         f"(PB = {T.PB_VALORES[e.nivel_protecao_spda.value]})")
    addP(f"<b>Sistema coordenado de DPS:</b> {T.PSPD_LABELS[e.pspd_nivel]} "
         f"(PSPD = {T.PSPD_VALORES[e.pspd_nivel]})")
    addP(f"<b>Medidas adicionais (PTA):</b> "
         f"{', '.join(T.PTA_LABELS[m] for m in e.pta_medidas)}")
    addS(10)

    # === 4. Linhas ===
    addP("4. Linhas elétricas conectadas", h1)
    if not projeto.linhas:
        addP("Nenhuma linha externa cadastrada.")
    else:
        dados = [["Nome", "Tipo", "Instalação", "Ambiente", "LL (m)", "UW (kV)"]]
        for linha in projeto.linhas:
            dados.append([
                linha.nome,
                linha.tipo.value.capitalize(),
                T.CI_LABELS[linha.instalacao.value][:25],
                T.CE_LABELS[linha.ambiente.value],
                f"{linha.LL:.0f}",
                f"{linha.UW}",
            ])
        elementos.append(tabela(dados, col_widths=[3.5 * cm, 2 * cm, 4 * cm, 2.5 * cm, 1.8 * cm, 2 * cm]))
    addS(10)

    # === 5. Resultados por zona ===
    addP("5. Componentes de risco por zona", h1)
    for rz in resultado.zonas:
        addP(rz.nome, h2)
        dados = [["Comp.", "Descrição", "N", "P", "L", "R (y⁻¹)"]]
        for nome, c in rz.componentes.items():
            dados.append([
                nome, c.descricao,
                f"{c.N:.2e}", f"{c.P:.2e}", f"{c.L:.2e}", f"{c.R:.2e}",
            ])
        elementos.append(tabela(dados, col_widths=[1.5 * cm, 5.8 * cm, 2 * cm, 2 * cm, 2 * cm, 2.4 * cm]))
        addS(4)
        addP(f"<b>R1 desta zona = {rz.R1:.3e} y⁻¹</b>")
        addS(8)

        # Frequência de danos F
        f = rz.frequencia
        addP(f"Frequência de danos F (Seção 7) — {rz.nome}", h2)
        crit_txt = "crítico (FT = 0,1/ano)" if f.sistema_critico else "não crítico (FT = 1/ano)"
        addP(f"<b>Classificação:</b> sistema {crit_txt}")

        cor_status = colors.HexColor("#c00000") if f.F > f.FT else colors.HexColor("#008000")
        dados_f = [
            ["Componente", "Fórmula (Tabela 7)", "Valor (1/ano)"],
            ["FB", "ND × PB (só ZPR0A)", f"{f.FB:.3e}"],
            ["FC", "ND × PC", f"{f.FC:.3e}"],
            ["FM", "NM × PM", f"{f.FM:.3e}"],
            ["FV", "(NL + NDJ) × PEB", f"{f.FV:.3e}"],
            ["FW", "(NL + NDJ) × PW", f"{f.FW:.3e}"],
            ["FZ", "NI × PZ", f"{f.FZ:.3e}"],
            ["F (total)", f"FT = {f.FT}/ano", f"{f.F:.3e}"],
        ]
        tF = tabela(dados_f, col_widths=[3 * cm, 6 * cm, 4 * cm])
        # Destaca linha total
        tF.setStyle(TableStyle([
            ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#fff2cc")),
            ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
            ("TEXTCOLOR", (2, -1), (2, -1), cor_status),
        ]))
        elementos.append(tF)
        addS(4)

        status_txt = (
            f"<b>Status:</b> F supera FT em "
            f"{(f.F/f.FT if f.FT > 0 else 0):.1f}× — necessita medidas adicionais."
            if f.F > f.FT else
            "<b>Status:</b> F ≤ FT — frequência de danos aceitável."
        )
        addP(status_txt)
        addS(10)

    # Quebra de página antes da conclusão
    elementos.append(PageBreak())

    # === 6. Conclusão ===
    addP("6. Riscos totais e conclusão", h1)
    dados = [["Risco", "Calculado", "Tolerável (RT)", "Status"]]
    riscos_linhas = [("R1 — Vida humana", resultado.R1, RT_R1, True)]
    if e.patrimonio_cultural:
        riscos_linhas.append(("R3 — Patrimônio cultural", resultado.R3, RT_R3, True))
    if projeto.avaliar_R4:
        riscos_linhas.append(("R4 — Perda econômica (Anexo D, informativo)", resultado.R4, RT_R4, False))

    for nome, R, RT, normativo in riscos_linhas:
        status = _status_texto(R, RT) if normativo else "—"
        dados.append([nome, f"{R:.3e}", f"{RT:.0e}", status])

    t = tabela(dados, col_widths=[7 * cm, 3 * cm, 3 * cm, 3 * cm])
    # Cor do status
    for i, (_, R, RT, normativo) in enumerate(riscos_linhas, start=1):
        if normativo:
            cor = colors.HexColor("#c00000") if R > RT else colors.HexColor("#008000")
            t.setStyle(TableStyle([
                ("TEXTCOLOR", (3, i), (3, i), cor),
                ("FONTNAME", (3, i), (3, i), "Helvetica-Bold"),
            ]))
    elementos.append(t)
    addS(12)

    # Diagnóstico
    if resultado.R1 > RT_R1:
        diagnostico = (
            f"<b>Conclusão:</b> O risco R1 calculado ({resultado.R1:.3e}) supera "
            f"o tolerável ({RT_R1:.0e}) em {resultado.R1/RT_R1:.1f}×. "
            "Medidas de proteção adicionais devem ser adotadas para reduzir R1 a RT."
        )
    else:
        diagnostico = (
            f"<b>Conclusão:</b> O risco R1 calculado ({resultado.R1:.3e}) está abaixo "
            f"do tolerável ({RT_R1:.0e}). A análise atende à NBR 5419-2:2026."
        )
    addP(diagnostico)
    addS(20)

    rodape = ParagraphStyle("rodape", parent=body, fontSize=8,
                            textColor=colors.grey, alignment=1)
    addP("Memorial gerado por Aplicativo de Análise de Risco SPDA — "
         "Vertec / NBR 5419-2:2026", rodape)

    doc.build(elementos)
    return buf.getvalue()
